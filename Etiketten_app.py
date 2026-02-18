import streamlit as st
import re
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import csv


def parse_ranges(range_str):
    parts = re.split(r",\s*", range_str)
    result = []
    for part in parts:
        if "-" in part:
            try:
                start, end = map(int, part.split("-"))
                result.extend([n for n in range(start, end + 1) if n > 0])
            except ValueError:
                pass
        else:
            try:
                n = int(part)
                if n > 0:
                    result.append(n)
            except ValueError:
                pass
    return result


def set_cell_spacing(paragraph, afstand_voor=None):
    p = paragraph._element
    pPr = p.get_or_add_pPr()

    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(afstand_voor) if afstand_voor is not None else "0")
    spacing.set(qn('w:after'), "0")
    spacing.set(qn('w:line'), "240")
    spacing.set(qn('w:lineRule'), "auto")
    pPr.append(spacing)

    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), "0")
    ind.set(qn('w:firstLine'), "0")
    pPr.append(ind)


def create_docx_table(labels):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.08)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0.39)
    section.right_margin = Inches(0.39)

    def add_label_table(label_block):
        table = doc.add_table(rows=9, cols=3)
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblCellMar = OxmlElement('w:tblCellMar')
        for side in ['left', 'right']:
            mar = OxmlElement(f'w:{side}')
            mar.set(qn('w:w'), '15')
            mar.set(qn('w:type'), 'dxa')
            tblCellMar.append(mar)
        for side in ['top', 'bottom']:
            mar = OxmlElement(f'w:{side}')
            mar.set(qn('w:w'), '0')
            mar.set(qn('w:type'), 'dxa')
            tblCellMar.append(mar)
        tblPr.append(tblCellMar)

        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        widths = [Inches(2.4805), Inches(2.4805), Inches(2.4805)]
        for row in table.rows:
            row.height = Inches(1.26)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            for i, cell in enumerate(row.cells):
                cell.width = widths[i]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if cell.paragraphs:
                    p = cell.paragraphs[0]._element
                    p.getparent().remove(p)

        idx = 0
        for row in table.rows:
            for cell in row.cells:
                if idx < len(label_block):
                    content = label_block[idx]
                    for i, line in enumerate(content):
                        if line.strip():
                            para = cell.add_paragraph()
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = para.add_run(line)
                            run.font.name = 'Arial'
                            run.font.size = Pt(10)

                            # Vette regels:
                            is_toegangsnummer = i == 1
                            is_inventarisregel = "Inventaris" in line or i == len(content) - 1

                            if is_toegangsnummer or is_inventarisregel:
                                run.bold = True
                                run.font.size = Pt(12)

                            para.paragraph_format.space_after = Pt(0)
                            para.paragraph_format.line_spacing_rule = 1
                            spacing_before = 226 if i == 0 else None
                            set_cell_spacing(para, afstand_voor=spacing_before)
                    idx += 1

    for i in range(0, len(labels), 27):
        label_block = labels[i:i + 27]
        add_label_table(label_block)
        if i + 27 < len(labels):
            doc.add_page_break()

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def split_title(titel, max_len=35):
    words = titel.split()
    lines = []
    current_line = ""
    for word in words:
        if len(current_line + " " + word) < max_len:
            current_line = (current_line + " " + word).strip()
        else:
            lines.append(current_line)
            current_line = word
    if current_line:
        lines.append(current_line)
    return lines[:2]  # maximaal 2 regels


def generate_box_labels(titel, nummer, groepen):
    labels = []
    for groep in groepen:
        if not groep.strip():
            continue
        gesplitste_titel = split_title(titel)
        labels.append([
            "Stadsarchief Amsterdam",
            f"{nummer}",
            *gesplitste_titel,
            groep.strip()
        ])
    return create_docx_table(labels)


def load_toegangstitels(csv_path="ToegangenLijst.csv"):
    mapping = {}
    try:
        with open(csv_path, newline='', encoding='utf-8-sig') as f:
            reader = csv.DictReader(f)
            for row in reader:
                nummer = row.get("toegangsnummer")
                titel = row.get("titel")
                if nummer and titel:
                    mapping[nummer.strip()] = titel.strip()
    except Exception as e:
        print("CSV leesfout:", e)
    return mapping


TOEGANGSTITELS = load_toegangstitels()


def main():
    st.title("📋 Stadsarchief Label Generator")

    option = st.radio("Wat wil je doen?", [
        "📁 Omslagetiketten maken",
        "📦 Doosetiketten maken"
    ])

    if option == "📁 Omslagetiketten maken":
        multi_ui = st.toggle("Meerdere toegang_
